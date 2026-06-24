"""
文档解析器 — DocumentParser

负责将 PDF / DOCX / Markdown / TXT 解析为结构化文本，
并返回包含元数据的 ParsedDocument 对象。

用法:
    parser = DocumentParser()
    doc = parser.parse("path/to/file.pdf", doc_type="textbook")
"""

from __future__ import annotations

import os
import pathlib
from typing import ClassVar, Literal

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from loguru import logger
from pydantic import BaseModel, Field


# ── Pydantic 模型 ──────────────────────────────────────────────────────────


class ParsedDocument(BaseModel):
    """文档解析结果。

    Attributes:
        text: 解析后的纯文本内容。
        filename: 不含路径的文件名（含扩展名）。
        page_count: PDF 的页数（非 PDF 文件为 0）。
        size_bytes: 文件大小（字节）。
        extension: 文件扩展名（小写，含点，如 ".pdf"）。
        scanned: PDF 是否为扫描件（无文字层）。
        doc_type: 文档用途类型 — "textbook"（教材）或 "past_paper"（真题）。
        course: 所属课程名（空=未分类），用于跨课程文档隔离。
    """

    text: str
    filename: str
    page_count: int = 0
    size_bytes: int = 0
    extension: str
    scanned: bool = False
    doc_type: Literal["textbook", "past_paper"] = "textbook"
    course: str = ""


# ── 文档解析器 ──────────────────────────────────────────────────────────────


class DocumentParser:
    """文档解析器（无状态，可复用）。

    职责:
        1. 前置校验 — 文件存在性、扩展名、大小。
        2. 按扩展名分派到具体解析器。
        3. 返回结构化的 ParsedDocument。

    边界情况处理（对应 TECH_DESIGN.md 第 9 节）:
        - ① 扫描版 PDF：提取后 text.strip()=="" → scanned=True
        - ② 加密 PDF：捕获加密异常 → raise PermissionError
        - ③ 文件 >200MB → raise ValueError
        - 扩展名不支持 → raise ValueError
        - 文件不存在 → FileNotFoundError
    """

    SUPPORTED_EXTENSIONS: ClassVar[set[str]] = {".pdf", ".docx", ".md", ".txt"}
    MAX_FILE_SIZE: ClassVar[int] = 200 * 1024 * 1024  # 200 MB

    def __init__(self) -> None:
        """初始化解析器。无状态，可复用。"""
        logger.debug("DocumentParser 初始化完成")

    # ── 入口 ────────────────────────────────────────────────────────────

    def parse(
        self,
        file_path: str | os.PathLike,
        doc_type: Literal["textbook", "past_paper"] = "textbook",
        course: str = "",
    ) -> ParsedDocument:
        """解析文档入口。

        执行流程:
            1. 将 file_path 转为 pathlib.Path 对象。
            2. _validate_file() — 前置校验。
            3. _dispatch() — 按扩展名分派解析。
            4. 组装 ParsedDocument 并返回。

        Args:
            file_path: 文件路径（str 或 PathLike）。
            doc_type: 文档用途类型。
            course: 所属课程名（空=未分类）。

        Returns:
            ParsedDocument: 包含解析文本和元数据。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 文件 >200MB / 扩展名不支持。
            PermissionError: PDF 已加密且无法解密。
        """
        path = pathlib.Path(file_path)
        logger.info("开始解析文档: {}", path.resolve())

        # 前置校验
        self._validate_file(path)

        # 获取文件基础信息
        size_bytes = path.stat().st_size
        extension = path.suffix.lower()

        # 分派解析
        text, extra_meta = self._dispatch(path)

        # 组装 ParsedDocument
        doc = ParsedDocument(
            text=text,
            filename=path.name,
            page_count=extra_meta.get("page_count", 0),
            size_bytes=size_bytes,
            extension=extension,
            scanned=extra_meta.get("scanned", False),
            doc_type=doc_type,
            course=course,
        )

        logger.info(
            "文档解析完成: {} | {} 字符 | {} 页 | 扫描件={}",
            doc.filename,
            len(doc.text),
            doc.page_count,
            doc.scanned,
        )
        return doc

    # ── 前置校验 ─────────────────────────────────────────────────────────

    def _validate_file(self, path: pathlib.Path) -> None:
        """前置校验：文件存在性、扩展名、大小。

        Args:
            path: 已转换为 pathlib.Path 的文件路径。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError: 扩展名不支持 或 文件 >200MB。
        """
        # ① 文件是否存在
        if not path.exists():
            msg = f"文件不存在: {path}"
            logger.error(msg)
            raise FileNotFoundError(msg)

        # ② 扩展名是否支持
        if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            msg = f"不支持的文件格式: {path.suffix}（支持: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}）"
            logger.warning(msg)
            raise ValueError(msg)

        # ③ 文件是否超过大小限制
        size = path.stat().st_size
        if size > self.MAX_FILE_SIZE:
            msg = (
                f"文件过大: {path.name}（{size / 1024 / 1024:.1f} MB），"
                f"请上传 {self.MAX_FILE_SIZE / 1024 / 1024:.0f} MB 以内的文件"
            )
            logger.warning(msg)
            raise ValueError(msg)

        logger.debug("前置校验通过: {} ({:.1f} KB)", path.name, size / 1024)

    # ── 扩展名分派 ───────────────────────────────────────────────────────

    def _dispatch(self, path: pathlib.Path) -> tuple[str, dict]:
        """按文件扩展名分派到具体解析器。

        Args:
            path: 文件路径（pathlib.Path）。

        Returns:
            (text, extra_metadata) — 解析文本和额外元数据字典。
            extra_metadata 包含 key: page_count, scanned（仅 PDF 有值）。
        """
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        elif ext in (".md", ".txt"):
            return self._parse_text(path)
        else:
            # 理论上不会跑到这里，_validate_file 已拦截
            msg = f"未预期的扩展名: {ext}"
            logger.error(msg)
            raise ValueError(msg)

    # ── PDF 解析 ─────────────────────────────────────────────────────────

    def _parse_pdf(self, path: pathlib.Path) -> tuple[str, dict]:
        """解析 PDF 文件。

        使用 PyMuPDF（fitz）逐页提取文本。

        加密检测:
            捕获 fitz.FileDataError 和 RuntimeError，
            将异常信息转为小写后检查是否包含 "password" 或 "encrypted"，
            以兼容不同 PyMuPDF 版本的异常类差异。

        扫描件检测:
            全部页面提取后 text.strip() == "" 则标记 scanned=True。

        Args:
            path: PDF 文件路径。

        Returns:
            (text, {"page_count": int, "scanned": bool})

        Raises:
            PermissionError: PDF 已加密且无法解密。
        """
        logger.debug("开始解析 PDF: {}", path.name)

        try:
            doc = fitz.open(str(path))
        except Exception as exc:
            # 兼容不同 PyMuPDF 版本的加密异常
            # 某些版本抛 fitz.FileDataError，某些版本抛 RuntimeError
            exc_msg = str(exc).lower()
            if "password" in exc_msg or "encrypted" in exc_msg:
                logger.warning("PDF 已加密: {}", path.name)
                raise PermissionError(f"PDF 已加密: {path.name}") from exc
            # 其他未知异常，原样抛出
            logger.opt(exception=True).error("打开 PDF 失败: {}", path.name)
            raise

        page_count = len(doc)
        pages_text: list[str] = []

        try:
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                pages_text.append(page_text)
        finally:
            doc.close()

        full_text = "\n".join(pages_text)

        # 扫描件检测：提取后无文字内容
        scanned = full_text.strip() == ""

        if scanned:
            logger.warning("检测到扫描件 PDF（无文字层）: {}", path.name)

        logger.debug(
            "PDF 解析完成: {} | {} 页 | {} 字符 | 扫描件={}",
            path.name,
            page_count,
            len(full_text),
            scanned,
        )
        return full_text, {"page_count": page_count, "scanned": scanned}

    # ── DOCX 解析 ────────────────────────────────────────────────────────

    def _parse_docx(self, path: pathlib.Path) -> tuple[str, dict]:
        """解析 .docx 文件。

        使用 python-docx 逐段落提取文本。
        注意：仅支持 .docx（Office 2007+），不支持旧版 .doc。

        Args:
            path: .docx 文件路径。

        Returns:
            (text, {"page_count": 0})
        """
        logger.debug("开始解析 DOCX: {}", path.name)

        try:
            doc = DocxDocument(str(path))
        except Exception:
            logger.opt(exception=True).error("打开 DOCX 失败: {}", path.name)
            raise

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                paragraphs.append(text)

        full_text = "\n".join(paragraphs)

        logger.debug(
            "DOCX 解析完成: {} | {} 段落 | {} 字符",
            path.name,
            len(paragraphs),
            len(full_text),
        )
        return full_text, {"page_count": 0, "scanned": False}

    # ── 文本文件解析（.md / .txt）───────────────────────────────────────

    def _parse_text(self, path: pathlib.Path) -> tuple[str, dict]:
        """解析 .md 或 .txt 文件。

        使用 UTF-8 编码读取全部内容。

        Args:
            path: .md 或 .txt 文件路径。

        Returns:
            (text, {"page_count": 0})
        """
        logger.debug("开始解析文本文件: {}", path.name)

        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 如果 UTF-8 解码失败，尝试用系统默认编码
            logger.warning("UTF-8 解码失败，尝试 UTF-8-sig: {}", path.name)
            text = path.read_text(encoding="utf-8-sig")
        except Exception:
            logger.opt(exception=True).error("读取文本文件失败: {}", path.name)
            raise

        logger.debug(
            "文本文件解析完成: {} | {} 字符",
            path.name,
            len(text),
        )
        return text, {"page_count": 0, "scanned": False}
